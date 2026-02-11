import io
import os
import re
import json
import zipfile
import tempfile
from pathlib import Path
from typing import Optional, List, Tuple, Dict, Any

import requests
import pandas as pd
import geopandas as gpd
from shapely.geometry import shape
from shapely.ops import unary_union
from pyproj import CRS

import matplotlib.pyplot as plt

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse, JSONResponse

from docx import Document
from docx.shared import Inches


# =========================
# CONFIG
# =========================
WFS_BASE = os.getenv("IDECOR_WFS_BASE", "https://idecor-ws.mapascordoba.gob.ar/geoserver/idecor/wfs")
DEFAULT_LAYERS = [
    "idecor:carta_suelo_50mil_2025",
    "idecor:carta_suelo_100mil_2025",
    "idecor:carta_suelo_250mil_2022",
    "idecor:carta_suelo_500mil_completa_2024",
]
DEFAULT_PROPS = (
    "simbolo,nombre,tipo_unida,cu,clase,subclase,ip,"
    "au_2m,au_1_5m,au_1m,"
    "serie_1,fase_1,porc_1,serie_2,fase_2,porc_2"
)

TIMEOUT = float(os.getenv("IDECOR_TIMEOUT", "40"))
USER_AGENT = os.getenv("IDECOR_UA", "soil-report-api/1.0")


app = FastAPI(title="Soil Report API (Córdoba - IDECOR)")


# =========================
# GEOMETRY + CRS HELPERS
# =========================
def utm_epsg_from_lonlat(lon: float, lat: float) -> int:
    zone = int((lon + 180) // 6) + 1
    return (32700 + zone) if lat < 0 else (32600 + zone)

def ensure_crs(gdf: gpd.GeoDataFrame) -> gpd.GeoDataFrame:
    if gdf.crs is None:
        # fallback razonable: si no hay CRS, asumir WGS84
        gdf = gdf.set_crs("EPSG:4326", allow_override=True)
    return gdf

def make_valid(gdf: gpd.GeoDataFrame) -> gpd.GeoDataFrame:
    gdf = gdf.copy()
    gdf["geometry"] = gdf["geometry"].buffer(0)
    gdf = gdf[gdf.geometry.notnull()].copy()
    return gdf

def dissolve_boundary(gdf: gpd.GeoDataFrame) -> gpd.GeoDataFrame:
    gdf = ensure_crs(gdf)
    gdf = make_valid(gdf)
    geom = unary_union(list(gdf.geometry))
    out = gpd.GeoDataFrame({"id": [1]}, geometry=[geom], crs=gdf.crs)
    return out

def bbox_expand(bounds: Tuple[float, float, float, float], margin_deg: float) -> Tuple[float, float, float, float]:
    minx, miny, maxx, maxy = bounds
    return (minx - margin_deg, miny - margin_deg, maxx + margin_deg, maxy + margin_deg)

def tiles_for_bbox(b: Tuple[float, float, float, float], tile_deg: float) -> List[Tuple[float, float, float, float]]:
    if tile_deg <= 0:
        return [b]
    minx, miny, maxx, maxy = b
    tiles = []
    x = minx
    while x < maxx:
        x2 = min(x + tile_deg, maxx)
        y = miny
        while y < maxy:
            y2 = min(y + tile_deg, maxy)
            tiles.append((x, y, x2, y2))
            y = y2
        x = x2
    return tiles


# =========================
# WFS HELPERS
# =========================
def wfs_get(params: Dict[str, Any]) -> requests.Response:
    headers = {"User-Agent": USER_AGENT, "Accept-Encoding": "gzip"}
    return requests.get(WFS_BASE, params=params, headers=headers, timeout=TIMEOUT)

def parse_hits_numberMatched(xml_text: str) -> int:
    m = re.search(r'numberMatched="(\d+)"', xml_text)
    return int(m.group(1)) if m else 0

def wfs_hits(typeNames: str, bbox4326: Tuple[float, float, float, float]) -> int:
    bbox_str = f"{bbox4326[0]},{bbox4326[1]},{bbox4326[2]},{bbox4326[3]},EPSG:4326"
    params = {
        "service": "WFS",
        "request": "GetFeature",
        "version": "2.0.0",
        "typeNames": typeNames,
        "resultType": "hits",
        "bbox": bbox_str,
    }
    r = wfs_get(params)
    if r.status_code != 200:
        raise HTTPException(status_code=502, detail=f"WFS hits error {r.status_code}: {r.text[:300]}")
    # hits suele devolver XML
    return parse_hits_numberMatched(r.text)

def wfs_results_page(
    typeNames: str,
    bbox4326: Tuple[float, float, float, float],
    propertyName: str,
    count: int,
    startIndex: int,
) -> Dict[str, Any]:
    bbox_str = f"{bbox4326[0]},{bbox4326[1]},{bbox4326[2]},{bbox4326[3]},EPSG:4326"
    params = {
        "service": "WFS",
        "request": "GetFeature",
        "version": "2.0.0",
        "typeNames": typeNames,
        "outputFormat": "application/json",
        "srsName": "EPSG:4326",
        "bbox": bbox_str,
        "resultType": "results",
        "count": count,
        "startIndex": startIndex,
    }
    if propertyName:
        params["propertyName"] = propertyName

    r = wfs_get(params)
    if r.status_code != 200:
        raise HTTPException(status_code=502, detail=f"WFS results error {r.status_code}: {r.text[:300]}")

    try:
        return r.json()
    except Exception:
        raise HTTPException(status_code=502, detail=f"WFS devolvió no-JSON en results: {r.text[:200]}")

def download_features_bbox(
    typeNames: str,
    bbox4326: Tuple[float, float, float, float],
    propertyName: str,
    count: int,
    max_pages: int,
) -> List[Dict[str, Any]]:
    feats: List[Dict[str, Any]] = []
    start = 0
    for _ in range(max_pages):
        data = wfs_results_page(typeNames, bbox4326, propertyName, count, start)
        page_feats = data.get("features", [])
        feats.extend(page_feats)
        returned = data.get("numberReturned", len(page_feats))
        if returned is None:
            returned = len(page_feats)
        if returned < count or len(page_feats) < count:
            break
        start += count
    return feats

def gdf_from_features(features: List[Dict[str, Any]]) -> gpd.GeoDataFrame:
    if not features:
        return gpd.GeoDataFrame(geometry=[], crs="EPSG:4326")

    geoms = []
    props = []
    for ft in features:
        geom = ft.get("geometry")
        if not geom:
            continue
        geoms.append(shape(geom))
        props.append(ft.get("properties", {}))

    gdf = gpd.GeoDataFrame(props, geometry=geoms, crs="EPSG:4326")
    gdf = gdf[gdf.geometry.notnull()].copy()
    return gdf


# =========================
# BUSINESS LOGIC
# =========================
def choose_layer_for_bbox(bbox4326: Tuple[float, float, float, float], layers: List[str]) -> Tuple[str, int]:
    for layer in layers:
        n = wfs_hits(layer, bbox4326)
        if n > 0:
            return layer, n
    return layers[-1], 0

def unit_key(row: pd.Series) -> str:
    simbolo = row.get("simbolo")
    nombre = row.get("nombre")
    if pd.notna(simbolo) and pd.notna(nombre):
        return f"{simbolo} — {nombre}"
    if pd.notna(simbolo):
        return str(simbolo)
    if pd.notna(nombre):
        return str(nombre)
    tipo = row.get("tipo_unida")
    return str(tipo) if pd.notna(tipo) else "Sin dato"

def save_map_units(gdf: gpd.GeoDataFrame, out_png: Path, title: str) -> None:
    if gdf.empty:
        return

    gdf = gdf.copy()
    gdf["unidad"] = gdf.apply(unit_key, axis=1)

    # “limpieza” de leyenda: top 12 + Otros
    top = gdf["unidad"].value_counts().head(12).index.tolist()
    gdf["unidad_plot"] = gdf["unidad"].where(gdf["unidad"].isin(top), "Otros")

    fig, ax = plt.subplots(figsize=(8.2, 6))
    gdf.plot(column="unidad_plot", ax=ax, legend=True)
    ax.set_title(title)
    ax.set_axis_off()
    fig.tight_layout()
    fig.savefig(out_png, dpi=180)
    plt.close(fig)

def save_map_ip(gdf: gpd.GeoDataFrame, out_png: Path, title: str) -> None:
    if gdf.empty or "ip" not in gdf.columns:
        return
    fig, ax = plt.subplots(figsize=(8.2, 6))
    gdf.plot(column="ip", ax=ax, legend=True)
    ax.set_title(title)
    ax.set_axis_off()
    fig.tight_layout()
    fig.savefig(out_png, dpi=180)
    plt.close(fig)

def make_docx(
    out_docx: Path,
    campo_area_ha: float,
    centroid_lonlat: Tuple[float, float],
    utm_epsg: int,
    layer_used: str,
    summary: pd.DataFrame,
    map_units: Optional[Path],
    map_ip: Optional[Path],
) -> None:
    doc = Document()
    doc.add_heading("Informe productivo de suelos (Córdoba - IDECOR)", level=0)

    lon, lat = centroid_lonlat
    doc.add_paragraph(f"Superficie del campo: {campo_area_ha:.1f} ha")
    doc.add_paragraph(f"Centroide (WGS84): lat {lat:.5f} / lon {lon:.5f}")
    doc.add_paragraph(f"UTM usada para áreas: EPSG:{utm_epsg}")
    doc.add_paragraph(f"Capa IDECOR utilizada: {layer_used}")

    doc.add_heading("Superficie por unidad / serie", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Unidad / Serie"
    hdr[1].text = "ha"
    hdr[2].text = "%"

    for _, r in summary.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["unidad"])
        row[1].text = f'{r["ha"]:.1f}'
        row[2].text = f'{r["pct"]:.1f}'

    if map_units and map_units.exists():
        doc.add_heading("Mapa de unidades / series", level=1)
        doc.add_picture(str(map_units), width=Inches(6.5))

    if map_ip and map_ip.exists():
        doc.add_heading("Mapa de índice de productividad (IP)", level=1)
        doc.add_picture(str(map_ip), width=Inches(6.5))

    doc.add_heading("Notas metodológicas (breve)", level=1)
    doc.add_paragraph(
        "Se tomó el límite del campo, se reproyectó a EPSG:4326 para consultar WFS IDECOR. "
        "La descarga se hizo por bbox con paginación y, si fue necesario, por tiles. "
        "Luego se realizó clip al polígono del campo y se calcularon áreas en UTM para obtener ha y %."
    )

    doc.save(out_docx)

def generate_report(
    boundary_gdf: gpd.GeoDataFrame,
    layers: List[str],
    margin_deg: float,
    tile_deg: float,
    count: int,
    max_pages: int,
    hit_threshold: int,
    propertyName: str,
) -> Tuple[bytes, str]:
    boundary_gdf = ensure_crs(boundary_gdf)
    boundary = dissolve_boundary(boundary_gdf).to_crs("EPSG:4326")
    minx, miny, maxx, maxy = boundary.total_bounds

    bbox0 = (minx, miny, maxx, maxy)
    bbox = bbox_expand(bbox0, margin_deg)

    centroid = boundary.geometry.iloc[0].centroid
    lon, lat = float(centroid.x), float(centroid.y)
    utm_epsg = utm_epsg_from_lonlat(lon, lat)

    layer_used, nmatch = choose_layer_for_bbox(bbox, layers)
    if nmatch == 0:
        raise HTTPException(status_code=404, detail="Sin hits IDECOR para el bbox del campo en todas las capas.")

    # Descargar con tiles si hace falta
    feats_all: List[Dict[str, Any]] = []
    tiles = tiles_for_bbox(bbox, tile_deg if nmatch > hit_threshold else 0)

    for tb in tiles:
        tn = wfs_hits(layer_used, tb)
        if tn == 0:
            continue
        feats = download_features_bbox(layer_used, tb, propertyName, count, max_pages)
        feats_all.extend(feats)

    gdf = gdf_from_features(feats_all)
    if gdf.empty:
        raise HTTPException(status_code=502, detail="No se pudieron bajar geometrías (results vacío).")

    # clip exacto
    gdf = make_valid(gdf)
    boundary = make_valid(boundary)
    gdf_clip = gpd.clip(gdf, boundary)

    if gdf_clip.empty:
        raise HTTPException(status_code=502, detail="Clip resultó vacío (revisar bbox/margen).")

    # áreas en UTM
    gdf_utm = gdf_clip.to_crs(CRS.from_epsg(utm_epsg))
    gdf_utm["ha"] = gdf_utm.geometry.area / 10000.0

    campo_utm = boundary.to_crs(CRS.from_epsg(utm_epsg))
    campo_area_ha = float(campo_utm.geometry.area.iloc[0] / 10000.0)

    # resumen
    gdf_utm["unidad"] = gdf_utm.apply(unit_key, axis=1)
    summary = (
        gdf_utm.groupby("unidad", dropna=False)["ha"]
        .sum()
        .reset_index()
    )
    summary["pct"] = (summary["ha"] / campo_area_ha) * 100.0
    summary = summary.sort_values("ha", ascending=False).reset_index(drop=True)

    # output temp (docx + mapas)
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        map_units = td / "map_units.png"
        map_ip = td / "map_ip.png"
        out_docx = td / "informe_suelos.docx"

        save_map_units(gdf_clip, map_units, "Unidades / Series (IDECOR)")
        save_map_ip(gdf_clip, map_ip, "Índice de Productividad (IP)")

        make_docx(
            out_docx=out_docx,
            campo_area_ha=campo_area_ha,
            centroid_lonlat=(lon, lat),
            utm_epsg=utm_epsg,
            layer_used=layer_used,
            summary=summary,
            map_units=map_units,
            map_ip=map_ip,
        )

        data = out_docx.read_bytes()

    filename = f"informe_suelos_{layer_used.replace(':','_')}.docx"
    return data, filename


# =========================
# LOAD BOUNDARY FILE
# =========================
def read_boundary_file(upload_path: Path) -> gpd.GeoDataFrame:
    suffix = upload_path.suffix.lower()

    if suffix in [".geojson", ".json"]:
        return gpd.read_file(upload_path)

    if suffix == ".kml":
        # requiere GDAL/Fiona con soporte KML
        return gpd.read_file(upload_path, driver="KML")

    if suffix == ".kmz":
        # kmz = zip con kml adentro
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            with zipfile.ZipFile(upload_path, "r") as z:
                z.extractall(td)
            kmls = list(td.rglob("*.kml"))
            if not kmls:
                raise HTTPException(status_code=422, detail="KMZ no contiene ningún .kml.")
            return gpd.read_file(kmls[0], driver="KML")

    if suffix == ".zip":
        # shapefile dentro
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            with zipfile.ZipFile(upload_path, "r") as z:
                z.extractall(td)
            shps = list(td.rglob("*.shp"))
            if not shps:
                raise HTTPException(status_code=422, detail="ZIP no contiene .shp (esperaba shapefile).")
            return gpd.read_file(shps[0])

    # Si suben shp suelto (no recomendado)
    if suffix == ".shp":
        return gpd.read_file(upload_path)

    raise HTTPException(status_code=422, detail=f"Formato no soportado: {suffix}. Subí ZIP(SHP) o KML/KMZ/GeoJSON.")


# =========================
# API ENDPOINTS
# =========================
@app.get("/health")
def health():
    return {"ok": True, "wfs_base": WFS_BASE}

@app.post("/soil/report")
async def soil_report(
    boundary_file: UploadFile = File(...),
    preferred_layer: str = Form("idecor:carta_suelo_50mil_2025"),
    tile_deg: float = Form(0.05),
    count: int = Form(50),
    max_pages: int = Form(50),
    hit_threshold: int = Form(800),
    margin_deg: float = Form(0.005),
    propertyName: str = Form(DEFAULT_PROPS),
):
    # layers: arranca por preferred, luego fallback
    layers = [preferred_layer] + [l for l in DEFAULT_LAYERS if l != preferred_layer]

    # guardar upload en temp
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        upath = td / boundary_file.filename
        content = await boundary_file.read()
        upath.write_bytes(content)

        boundary_gdf = read_boundary_file(upath)
        boundary_gdf = ensure_crs(boundary_gdf)

        docx_bytes, filename = generate_report(
            boundary_gdf=boundary_gdf,
            layers=layers,
            margin_deg=margin_deg,
            tile_deg=tile_deg,
            count=count,
            max_pages=max_pages,
            hit_threshold=hit_threshold,
            propertyName=propertyName,
        )

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
